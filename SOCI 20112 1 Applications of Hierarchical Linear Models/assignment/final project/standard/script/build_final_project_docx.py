"""Build revised final project Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "final project_revised.docx"
FIG = BASE / "output" / "figures" / "propensity_score_histogram.png"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Final Project: Sports Participation and First-Grade Reading Achievement")
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph()

# Introduction
add_heading(doc, "Introduction", 1)
add_para(
    doc,
    "This project examines whether participation in organized sports is associated with "
    "first-grade reading achievement using data from the Early Childhood Longitudinal Study—"
    "Kindergarten Class of 2010–11 (ECLS-K:2011). Although sports participation is often "
    "linked to academic success, the causal effect is unclear because participation is "
    "selective: more advantaged students and stronger readers may be more likely to join "
    "sports programs, and some school programs require minimum grades to participate. "
    "I address this question using two approaches required for the project: (1) covariance "
    "adjustment with hierarchical linear models (HLM), and (2) propensity score stratification "
    "followed by a final HLM in HLM software."
)

# Data
add_heading(doc, "Data and Variables", 1)
add_para(
    doc,
    "The analytic sample includes 12,436 students nested in 817 schools. The outcome variable "
    "is Read1, a first-grade reading test score (M = 78.91, SD = 23.74). The treatment variable "
    "is Sports, a binary indicator of organized sports participation (37.1% participated). "
    "Student-level covariates include gender (Male), racial minority status (NonWhite), a "
    "composite SES measure, arts participation (Arts), kindergarten reading score (ReadK), "
    "and age in months (Age). School-level covariates include Catholic school (Catholic), "
    "other private school (OthPriv), and kindergarten enrollment (SchKEnrl), with public "
    "schools as the reference group."
)
add_table(
    doc,
    ["N", "Schools", "Mean Read1", "SD Read1", "Sports Rate"],
    [["12,436", "817", "78.91", "23.74", "37.1%"]],
)

# Analytic approach
add_heading(doc, "Analytic Approach", 1)
add_para(
    doc,
    "Because students are nested within schools, I estimated two-level hierarchical linear "
    "models with students at Level 1 and schools at Level 2. All covariance-adjustment models "
    "were estimated in R using lme4/lmerTest. The final propensity score stratified model was "
    "estimated in HLM 7 using the PS-stratified Level-1 file (eclsk_ps_hlm.sav) linked to "
    "schools by S1_ID."
)

# Section I
add_heading(doc, "I. Covariance Adjustment", 1)

add_heading(doc, "A. Naive Model (No Covariates)", 2)
add_para(doc, "Model specification:", bold=True)
add_para(
    doc,
    "Level-1: Read1 = π0 + π1(Sports) + e\n"
    "Level-2: π0 = β00 + u0; π1 = β10 + u1\n"
    "Combined: Read1 = β00 + β10(Sports) + u0 + u1(Sports) + e\n\n"
    "Only Sports is included as a predictor, and both the intercept and Sports slope are allowed "
    "to vary across schools: (1 + Sports | S1_ID)."
)
add_table(
    doc,
    ["Effect", "Estimate", "SE", "t", "p"],
    [
        ["Intercept", "77.23", "0.44", "174.39", "< .001"],
        ["Sports", "4.04", "0.44", "9.10", "< .001"],
    ],
)
add_para(doc, "Random effects:", bold=True)
add_table(
    doc,
    ["Component", "Variance", "SD"],
    [
        ["School intercept", "106.2", "10.31"],
        ["School Sports slope", "11.5", "3.39"],
        ["Residual", "455.7", "21.35"],
    ],
)
add_para(
    doc,
    "Interpretation. The intercept of 77.23 is the estimated average first-grade reading score "
    "for non-participants (Sports = 0) in an average school. The Sports coefficient of 4.04 "
    "indicates that, in the naive model, sports participants scored about 4.04 points higher "
    "than non-participants, holding the school context at its average (p < .001). "
    "The significant school intercept variance (SD = 10.31) shows that average reading "
    "achievement differed substantially across schools. The random slope variance for Sports "
    "was positive but did not reach conventional significance at α = .05, suggesting limited "
    "evidence that the sports–reading association varied across schools. "
    "In plain language, the naive model suggests a sizable sports advantage, but this estimate "
    "is likely inflated by selection because it ignores student and school differences."
)

add_heading(doc, "B. Covariate-Adjusted Model", 2)
add_para(doc, "Model specification:", bold=True)
add_para(
    doc,
    "Level-1: Read1 = π0 + π1(Sports) + π2(Male) + … + π9(SchKEnrl) + e\n"
    "Level-2: π0 = β00 + u0; all π coefficients for Level-1 predictors are fixed (no random slopes)\n"
    "Combined: Read1 = β00 + β10(Sports) + … + β90(SchKEnrl) + u0 + e\n\n"
    "All student- and school-level covariates are included with fixed slopes. Only the school "
    "intercept is random: (1 | S1_ID)."
)
add_table(
    doc,
    ["Predictor", "Estimate", "SE", "t", "p"],
    [
        ["Sports", "0.97", "0.30", "3.22", ".001"],
        ["Male", "−1.15", "0.27", "−4.31", "< .001"],
        ["NonWhite", "−1.02", "0.32", "−3.14", ".002"],
        ["SES", "2.53", "0.21", "12.21", "< .001"],
        ["Arts", "1.07", "0.39", "2.75", ".006"],
        ["ReadK", "1.23", "0.01", "118.91", "< .001"],
        ["Age", "−0.05", "0.03", "−1.72", ".086"],
        ["Catholic", "−0.04", "0.68", "−0.06", ".950"],
        ["OthPriv", "1.03", "0.75", "1.37", ".170"],
        ["SchKEnrl", "−0.00", "0.01", "−0.74", ".458"],
    ],
)
add_para(
    doc,
    "Interpretation of the treatment effect. After adjusting for observed covariates, the Sports "
    "coefficient fell from 4.04 to 0.97 (p = .001). Sports participants scored about one point "
    "higher in first-grade reading than otherwise similar non-participants."
)
add_para(
    doc,
    "Why the treatment effect changed. The large drop from 4.04 to 0.97 suggests that much of "
    "the naive association was due to observed confounding. Students with higher SES and stronger "
    "kindergarten reading scores were more likely both to participate in sports and to score "
    "higher on Read1. Once these differences were controlled, the sports effect was much smaller."
)
add_para(
    doc,
    "What else I learned from the model. Kindergarten reading (ReadK) was the strongest predictor "
    "of first-grade reading (b = 1.23, p < .001). SES was also strongly positive (b = 2.53). "
    "Male students and racial minority students scored lower on average after adjustment. Arts "
    "participation was positively associated with reading (b = 1.07). School sector and enrollment "
    "were not statistically significant once student covariates were included."
)

# Section II
add_heading(doc, "II. Propensity Score Stratification", 1)
add_para(
    doc,
    "The goal of propensity score stratification is to estimate the sports–reading association "
    "within groups of students who have the same, or nearly the same, probability of sports "
    "participation."
)

add_heading(doc, "A. Building the Propensity Score Model", 2)
add_para(doc, "Model specification:", bold=True)
add_para(
    doc,
    "Level-1: logit(P(Sports = 1)) = η\n"
    "η = γ00 + γ10(Male) + γ20(NonWhite) + … + γ90(SchKEnrl)\n"
    "Level-2: γ00 = δ00 + v0 (random intercept for school)\n\n"
    "I estimated a two-level logistic regression model predicting Sports from all pre-treatment "
    "covariates, with a random intercept for school (S1_ID). Predicted log odds of treatment "
    "(the linear predictor from the logistic model) were saved for each student and converted "
    "to propensity scores (predicted probabilities)."
)
add_para(doc, "Selected fixed effects from the two-level logistic model:", bold=True)
add_table(
    doc,
    ["Predictor", "Log-odds coefficient"],
    [
        ["Male", "0.93"],
        ["NonWhite", "−0.89"],
        ["SES", "0.63"],
        ["Arts", "0.25"],
        ["ReadK", "0.006"],
        ["Age", "0.023"],
        ["Catholic", "0.37"],
        ["OthPriv", "−0.27"],
        ["SchKEnrl", "0.0003 (n.s.)"],
    ],
)
add_para(
    doc,
    "Male, NonWhite, SES, Arts, ReadK, Age, Catholic, and OthPriv were significant predictors "
    "of sports participation. These results confirm that treatment is not random and justify "
    "propensity score adjustment."
)

add_heading(doc, "B. Common Support", 2)
add_para(
    doc,
    "I examined the distribution of propensity scores for treated (Sports = 1) and control "
    "(Sports = 0) students. Although sports participants tended to have higher propensity scores, "
    "there was substantial overlap between groups. To ensure common support, I removed students "
    "with propensity scores below 0.05 or above 0.95. This removed 48 students (0.4% of the "
    "sample), leaving 12,388 students for stratification."
)
add_para(
    doc,
    "Why this step is necessary. Without common support, treated and control students at extreme "
    "propensity scores are not comparable: there may be treated students with very high propensity "
    "scores but no similar control students (or vice versa). Estimating treatment effects in "
    "those regions requires extrapolation and can produce biased estimates. Restricting the sample "
    "to the region of overlap ensures comparisons are made among students with similar "
    "probabilities of treatment."
)
if FIG.exists():
    doc.add_paragraph("Figure 1. Propensity score distribution by sports participation status.")
    doc.add_picture(str(FIG), width=Inches(5.5))
    doc.add_paragraph()

add_heading(doc, "C. Stratification", 2)
add_para(
    doc,
    "Students in the common-support sample were sorted by propensity score and divided into "
    "five equal-sized strata using quintiles. Stratum sizes were approximately 2,477–2,478 "
    "students per stratum. Stratum 1 served as the reference category in the final model; "
    "strata 2–5 were entered as dummy indicators (strata2–strata5)."
)

add_heading(doc, "D. Balance Diagnostics", 2)
add_para(
    doc,
    "I assessed balance between treated and control students using absolute standardized mean "
    "differences (|SMD|) on all covariates, both overall and within each stratum."
)
add_table(
    doc,
    ["Balance criterion", "Result"],
    [
        ["Within-stratum |SMD| below 0.10", "80.0%"],
        ["Within-stratum |SMD| below 0.20", "88.9%"],
    ],
)
add_para(
    doc,
    "Quality of balance. After stratification, 80.0% of within-stratum comparisons had |SMD| "
    "below 0.10 and 88.9% below 0.20, indicating good balance by conventional standards "
    "(|SMD| < 0.10 is ideal; |SMD| < 0.20 is acceptable). This means that, within each stratum, "
    "treated and control students were much more similar on observed covariates than in the "
    "full sample. Improved balance reduces bias from observed confounding, although it does not "
    "eliminate bias from unmeasured confounders."
)

add_heading(doc, "E. Final Analytic Model (HLM)", 2)
add_para(doc, "Model specification:", bold=True)
add_para(
    doc,
    "Level-1: Read1 = π0 + π1(Sports) + π2(strata2) + … + π5(strata5) + e\n"
    "Level-2: π0 = β00 + u0 (random intercept only)\n\n"
    "This model was estimated in HLM 7 using the PS-stratified Level-1 file linked by S1_ID."
)
add_table(
    doc,
    ["Effect", "Estimate", "SE", "t", "p"],
    [
        ["Intercept (Stratum 1)", "70.81", "0.55", "127.84", "< .001"],
        ["Sports", "1.08", "0.44", "2.45", ".014"],
        ["Stratum 2", "4.74", "0.63", "7.48", "< .001"],
        ["Stratum 3", "7.85", "0.66", "11.93", "< .001"],
        ["Stratum 4", "9.76", "0.68", "14.37", "< .001"],
        ["Stratum 5", "16.14", "0.72", "22.31", "< .001"],
    ],
)
add_para(
    doc,
    "Interpretation. Within propensity score strata, sports participants scored about 1.08 points "
    "higher on first-grade reading than comparable non-participants (p = .014). This estimate "
    "is much smaller than the naive estimate (4.04) and similar to the covariate-adjusted estimate "
    "(0.97), suggesting that most of the raw difference was due to selection on observed "
    "characteristics. The significant stratum coefficients reflect higher average reading scores "
    "in higher propensity-score strata, which is expected because students with higher "
    "propensity scores tend to have higher SES and kindergarten achievement."
)

add_heading(doc, "E. Causal Interpretation", 2)
add_para(
    doc,
    "A causal interpretation of the sports effect would require several assumptions: (1) no "
    "unmeasured confounding— all variables that affect both sports participation and reading "
    "are observed and correctly modeled; (2) common support— treated and control units have "
    "overlapping propensity scores; (3) correct model specification; and (4) stable unit "
    "treatment value (SUTVA)— one student's sports participation does not affect another "
    "student's reading."
)
add_para(
    doc,
    "Assessment. Common support and balance on observed covariates appear reasonable, so bias "
    "from observed confounding is likely reduced. However, the no-unmeasured-confounding "
    "assumption is questionable. Factors such as parental involvement, student motivation, "
    "neighborhood resources, health, and school encouragement of extracurricular activities may "
    "affect both sports participation and reading but were not measured. Because sports programs "
    "often require minimum grades, reverse causality is also possible. Therefore, the results "
    "should be interpreted as evidence of an association rather than definitive proof of a "
    "causal effect."
)

add_heading(doc, "F. Sensitivity Analysis", 2)
add_para(
    doc,
    "Following the course sensitivity approach, I evaluated whether the estimated sports effect "
    "is robust to potential unobserved confounding. I selected kindergarten reading (ReadK) as "
    "a powerful predictor of the outcome (b = 1.23 in the adjusted outcome model) and computed "
    "the mean difference in ReadK between sports participants and non-participants "
    "(Δ = 3.55 points). The estimated plausible bias from an unobserved confounder as strong "
    "as ReadK is:"
)
add_para(doc, "Plausible bias = |bReadK| × |ΔReadK| = 1.23 × 3.55 ≈ 4.37 points.", bold=True)
add_para(
    doc,
    "The PS-stratified sports coefficient (1.08) and covariate-adjusted coefficient (0.97) are "
    "both smaller than this plausible bias (4.37). This means that if an unobserved confounder "
    "as strongly related to reading as kindergarten reading score were omitted from the model, "
    "it could explain the remaining estimated sports effect. The naive estimate (4.04) is "
    "close to this benchmark, which is consistent with the view that much of the raw association "
    "was confounding. Thus, the small adjusted effects are not robust to unobserved confounding "
    "of ReadK-like strength, and causal claims should be made cautiously."
)
add_para(doc, "Model specification sensitivity (comparison across approaches):", bold=True)
add_table(
    doc,
    ["Model", "Sports coefficient", "SE", "t", "95% CI"],
    [
        ["Naive HLM", "4.04", "0.44", "9.10", "[3.17, 4.91]"],
        ["Covariate-adjusted HLM", "0.97", "0.30", "3.22", "[0.38, 1.56]"],
        ["PS-stratified HLM", "1.08", "0.44", "2.45", "[0.22, 1.95]"],
        ["PS-stratified + adjusted HLM", "0.96", "0.30", "3.24", "[0.38, 1.54]"],
    ],
)
add_para(
    doc,
    "The three adjusted estimates (0.97, 1.08, and 0.96) are very similar, suggesting that "
    "the conclusion of a small positive association is moderately robust to model specification. "
    "The large difference between the naive and adjusted estimates confirms that confounding "
    "was substantively important."
)

# Section III
add_heading(doc, "III. Conclusion", 1)
add_para(
    doc,
    "Overall, sports participation was positively associated with first-grade reading achievement. "
    "The naive HLM suggested a large 4-point advantage for sports participants, but this "
    "estimate fell to about 1 point after covariate adjustment and propensity score "
    "stratification. This pattern indicates that much of the raw difference in reading "
    "achievement reflected pre-existing differences between students rather than sports "
    "participation itself."
)
add_para(
    doc,
    "Comparison of methods. Covariance adjustment directly controls for observed confounders in "
    "the outcome model and is straightforward to implement in HLM. Its main advantage is "
    "efficiency: it uses the full sample and provides interpretable coefficients for all "
    "covariates. Its limitation is that it depends heavily on the correctness of the outcome "
    "model and does not explicitly evaluate balance between treated and control groups."
)
add_para(
    doc,
    "Propensity score stratification improves comparability by grouping students with similar "
    "treatment probabilities and checking balance with SMDs. Its advantage is transparency "
    "about selection and balance; the stratification step makes it clear whether treated and "
    "control students are comparable on observed covariates. Its limitations include sample "
    "loss from common-support trimming, reliance on the propensity model, and the need to "
    "model strata in the final outcome analysis."
)
add_para(
    doc,
    "Both methods converged on a similar conclusion: a small but statistically significant "
    "positive association between sports and reading after addressing observed selection. "
    "However, sensitivity analysis showed that the adjusted effects are not robust to "
    "unobserved confounding as strong as kindergarten reading achievement. Future research "
    "with richer measures of family and student characteristics would be needed to strengthen "
    "causal inference."
)

doc.save(OUT)
print(f"Saved: {OUT}")
